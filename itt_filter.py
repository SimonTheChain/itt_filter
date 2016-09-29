#!/usr/bin/python
# -*- coding: utf-8 -*-

# itt filter
#
# Author: Simon Lachaîne


import os
import sys
import time

from PyQt4 import QtCore, QtGui
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

import itt_filter_ui as main_frame


itt_path = ""
destination_path = ""


class FilterItt(QtCore.QThread):

    def __init__(self):
        QtCore.QThread.__init__(self)

    def __del__(self):
        self.wait()

    def run(self):
        try:
            approot = os.path.dirname(os.path.abspath(__file__))
        except NameError:
            approot = os.path.dirname(os.path.abspath(sys.argv[0]))

        os.chdir(approot)
        document = Document("temp.docx")

        with open(itt_path) as itt:

            for line in itt:
                clean_line = line.lstrip().decode("utf-8")

                if clean_line.startswith("<p"):

                    if "italic" in clean_line:
                        head, sep, tail = clean_line.partition("begin")
                        final_line, sep, tail = (sep + tail).rpartition("</p>")
                        paragraph = document.add_paragraph(
                            final_line.replace(
                                "<br/>", " ").replace(
                                "&apos;", "'").replace(
                                "span ", "").replace(
                                "</span>", ""
                            ))
                        paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.TURQUOISE

                    if '"' in clean_line.split(">")[1] and "italic" not in clean_line:
                        head, sep, tail = clean_line.partition("begin")
                        final_line, sep, tail = (sep + tail).rpartition("</p>")
                        paragraph = document.add_paragraph(
                            final_line.replace(
                                "<br/>", " ").replace(
                                "&apos;", "'").replace(
                                "span ", "").replace(
                                "</span>", ""
                            ))
                        paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

                    if '<<' in clean_line.split(">")[1]:
                        head, sep, tail = clean_line.partition("begin")
                        final_line, sep, tail = (sep + tail).rpartition("</p>")
                        paragraph = document.add_paragraph(
                            final_line.replace(
                                "<br/>", " ").replace(
                                "&apos;", "'").replace(
                                "span ", "").replace(
                                "</span>", ""
                            ))
                        paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

        filtered = os.path.splitext(itt_path)[0]
        os.chdir(destination_path)

        document.save("%s.docx" % os.path.basename(filtered))


class IttFilterApp(QtGui.QMainWindow, main_frame.Ui_FilterWindow):
    def __init__(self, parent=None):
        super(IttFilterApp, self).__init__(parent)

        self.default_directory = ""
        self.filter_itt = FilterItt()

        self.setupUi(self)

        self.itt_btn.clicked.connect(self.itt_dlg)
        self.destination_btn.clicked.connect(self.destination_dlg)
        self.filter_btn.clicked.connect(self.filter)

        self.show()

    def itt_dlg(self):
        global itt_path
        itt_path = str(QtGui.QFileDialog.getOpenFileName(
                parent=self,
                caption="Locate subtitle file",
                directory=self.default_directory,
                filter="Itt files (*.itt)"))

        self.itt_lbl.setText(os.path.basename(itt_path))

        if itt_path:
            self.default_directory = os.path.dirname(itt_path)

    def destination_dlg(self):
        global destination_path
        destination_path = str(QtGui.QFileDialog.getExistingDirectory(
            parent=self,
            caption="Choose destination folder",
            directory=self.default_directory))

        self.destination_lbl.setText(os.path.basename(destination_path))

        if destination_path:
            self.default_directory = destination_path

    def filter(self):
        if itt_path == "":
            itt_msg = QtGui.QMessageBox()

            itt_msg.setIcon(QtGui.QMessageBox.Information)
            itt_msg.setText("Please locate the itt file to process.")
            itt_msg.setWindowTitle("Input needed")
            itt_msg.setStandardButtons(QtGui.QMessageBox.Ok)
            itt_msg.exec_()
            return

        if destination_path == "":
            dest_msg = QtGui.QMessageBox()

            dest_msg.setIcon(QtGui.QMessageBox.Information)
            dest_msg.setText("Please select the destination folder.")
            dest_msg.setWindowTitle("Input needed")
            dest_msg.setStandardButtons(QtGui.QMessageBox.Ok)
            dest_msg.exec_()
            return

        else:
            self.filter_btn.setEnabled(False)
            self.results_lbl.setText("Filtering...")
            self.connect(self.filter_itt, QtCore.SIGNAL("finished()"), self.filter_done)
            self.filter_itt.start()

    def filter_done(self):
        self.results_lbl.setText("Filtering finished at %s." % time.strftime("%X"))
        self.filter_btn.setEnabled(True)


def main():
    app = QtGui.QApplication(sys.argv)
    gui = IttFilterApp()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
